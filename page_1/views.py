from django.contrib import messages
from django.contrib.auth import login, logout
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.views import LoginView
from django.http import HttpResponse

from django.shortcuts import render,reverse
from django.views.generic import TemplateView,ListView, View
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from .models import UyRasmlari, Uylar, Xususiyati, Vazifalar,Review

from django.shortcuts import render, redirect
from .forms import MijozForm, UylarForm, XususiyatForm, UypictureForm, UserLoginForm, UserRegistrForm,ReviewForm


class Home_sahifasi(View):
    def get(self, request):
        xususiyatlar = Xususiyati.objects.all()
        vazifa = Vazifalar.objects.all()


        return render(request, "home.html", {"xususiyatlar":xususiyatlar, "vaz":vazifa})


class AdminCheck(UserPassesTestMixin):
    def test_func(self):
        return self.request.user.is_superuser


class mijozqoshish(View):
    def get(self,request):
        form=MijozForm()
        return render(request,"mijozqoshish.html",{"form":form})
    def post(self,request):
        form=MijozForm(data=request.POST)
        if form.is_valid():
            form.save()
            return redirect('uylar')
        else:
            return render(request,"mijozqoshish.html",{"form":form})



class xususiyat(View):
    def get(self,request):
        form=XususiyatForm()
        return render(request,"xususiyatlari.html",{"form":form})
    def post(self,request):
        form=XususiyatForm(data=request.POST,files=request.FILES)
        if form.is_valid():
            form.save()
            return redirect('home')
        else:
            return render(request,"xususiyatlari.html",{"form":form})


class uylar(View):
    def get(self,request):
        form=UylarForm()
        return render(request,"uylar.html",{"form":form})
    def post(self,request):
        form=UylarForm(data=request.POST,files=request.FILES)
        if form.is_valid():
            form.save()
            return redirect('xususiyat')
        else:
            return render(request,"uylar.html",{"form":form})




class detail_home(View):
    def get(self, request, a):
        review_form = ReviewForm()
        uy = Xususiyati.objects.get(id=a).uy
        # print(uy)
        sharxlar = Review.objects.filter(uy_id=uy.id)
        xususiyat = Xususiyati.objects.get(pk=a)
        Uy_picture = UyRasmlari.objects.filter(xususiyat_id=a)
        return render(request, 'detail_img.html', {'Uy_picture':Uy_picture, "xususiyat":xususiyat,'form': review_form,'sharxlar': sharxlar})

    def post(self, request, a):
        review_form = ReviewForm(data=request.POST)
        uy = Xususiyati.objects.get(id=a).uy
        # print(uy)
        sharxlar = Review.objects.filter(uy_id=uy.id)
        # print(sharxlar)
        if review_form.is_valid():
            Review.objects.create(
                uy_id=uy,
                user_id=request.user,
                description=review_form.cleaned_data['description'],
                stars=review_form.cleaned_data["stars"]
            )
            return redirect(reverse("detail_home", kwargs={'a': uy.id}))
        print(sharxlar)
        return render(request, "detail_img.html", {'uy': uy, 'sharxlar': sharxlar, 'form': review_form})
from docx import *
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.views import View
from .models import Uylar, Xususiyati
from docx import Document
from docx.shared import Inches

class DownloadWordView(View):
    def get(self, request, pk, *args, **kwargs):
        # Uylar modelidan pk orqali uy obyektini olish

        # Ushbu uyga tegishli Xususiyati obyektlarini olish
        xususiyatlar = Xususiyati.objects.filter(id=pk)
        xususiyatlar2 = Xususiyati.objects.get(id=pk)
        # Yangi Word hujjati yaratish
        document = Document()
        document.add_heading('Shartnoma', level=0)

        # Uy egasining ma'lumotlarini hujjatga qo'shish
        document.add_paragraph(f'Uy egasi ismi: {xususiyatlar2.uy.mijoz_id.ismi}')
        document.add_paragraph(f'Uy egasi familiyasi: {xususiyatlar2.uy.mijoz_id.familyasi}')
        document.add_paragraph(f'Uy egasi telefon raqami: {xususiyatlar2.uy.mijoz_id.tel}')

        # Xususiyatlar jadvalini yaratish
        table = document.add_table(rows=1, cols=8)
        table.style = 'Table Grid'  # Jadval stilini o'rnatish

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Holati'
        hdr_cells[1].text = 'Yashash Maydoni'
        hdr_cells[2].text = 'Xonalar Soni'
        hdr_cells[3].text = 'Joylashuvi'
        hdr_cells[4].text = 'Vazifasi'
        hdr_cells[5].text = 'Narxi'
        hdr_cells[6].text = 'Sana'
        hdr_cells[7].text = 'Star'

        for xususiyat in xususiyatlar:
            row_cells = table.add_row().cells
            row_cells[0].text = str(xususiyat.uy.holati)
            row_cells[1].text = str(xususiyat.uy.yashash_maydoni)
            row_cells[2].text = str(xususiyat.uy.xonalar_soni)
            row_cells[3].text = str(xususiyat.uy.joylashuvi)
            row_cells[4].text = str(xususiyat.vazifasi.vazifa)
            row_cells[5].text = str(xususiyat.narxi)
            row_cells[6].text = xususiyat.sana.strftime("%Y-%m-%d %H:%M:%S")
            row_cells[7].text = str(xususiyat.star.id)

        # Hujjatni yozish uchun HttpResponse tayyorlash
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=details_of_{xususiyatlar2.uy.id}.docx'
        document.save(response)
        return response
class addpicture(View):
    def get(self,request):
        form=UypictureForm()
        return render(request,"addpicturre.html",{"form":form})
    def post(self,request):
        form=UypictureForm(data=request.POST,files=request.FILES)
        if form.is_valid():
            form.save()
            return redirect('addpicture')
        else:
            return render(request,"addpicturre.html",{"form":form})


class LoginView(View):
    def get(self, request):
        user = UserLoginForm()
        return render(request, "login.html", {'user':user})

    def post(self, request):
        check = AuthenticationForm(data=request.POST)

        if check.is_valid():
            user = check.get_user()
            login(request, user)
            messages.success(request, "Siz tizimga kirdingiz!")
            return redirect('home')
        else:
            return redirect('login')

class LogoutView(View):
    def get(self, request):
        logout(request)
        messages.info(request, "Siz tizimdan chiqdingiz!")
        return redirect("home")


class HomesView(View):
    def get(self, request,id):
        uylar = Xususiyati.objects.filter(vazifasi=id)
        xususiyatlar = Xususiyati.objects.all()
        vazifa = Vazifalar.objects.all()

        search_query = request.GET.get('q','')
        if search_query:
            uylar = Xususiyati.objects.filter(uy__joylashuvi__icontains=search_query)




        return render(request, "home.html", {"xususiyatlar": xususiyatlar, "vaz": vazifa,"uylar":uylar,"search":search_query})

class UserRegistrView(View):
    def get(self,request):
        form=UserRegistrForm()
        return render(request,'register.html',{'form':form})

    def post(self, request):
        form = UserRegistrForm(data=request.POST)
        if form.is_valid():
            form.save()
            return redirect("login")
        else:
            return render(request, "register.html", {'form': form})